import os
import PyPDF2
import docx
from pathlib import Path
import logging
from .utils import clean_text

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor:
    """Handle file upload and text extraction from various formats"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_from_pdf,
            'text/plain': self._extract_from_txt,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_from_docx,
            'application/msword': self._extract_from_docx
        }
    
    def extract_text(self, file_path, file_type):
        """Extract text from uploaded file based on type"""
        try:
            if file_type in self.supported_formats:
                text = self.supported_formats[file_type](file_path)
                return clean_text(text)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF files using PyPDF2"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            return ""
        
        return text
    
    def _extract_from_txt(self, file_path):
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {str(e)}")
                return ""
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text += paragraph.text + "\n"
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {str(e)}")
            return ""
    
    def get_page_count(self, file_path, file_type):
        """Get approximate page count for summary length calculation"""
        try:
            if file_type == 'application/pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    return len(pdf_reader.pages)
            else:
                # Estimate pages based on text length (approximately 250 words per page)
                text = self.extract_text(file_path, file_type)
                word_count = len(text.split())
                return max(1, word_count // 250)
        except:
            return 1  # Default to 1 page if unable to determine
    
    def validate_file(self, file_path, max_size_mb=10):
        """Validate uploaded file"""
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return False, "File not found"
            
            # Check file size
            file_size = os.path.getsize(file_path) / (1024 * 1024)  # Convert to MB
            if file_size > max_size_mb:
                return False, f"File size ({file_size:.1f}MB) exceeds limit ({max_size_mb}MB)"
            
            return True, "File is valid"
            
        except Exception as e:
            return False, f"File validation error: {str(e)}"